import json
import os
from io import BytesIO
from pathlib import Path

import requests
from flask import Flask, flash, render_template, request


OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
OLLAMA_TIMEOUT = int(os.environ.get("OLLAMA_TIMEOUT", "300"))
MAX_UPLOAD_BYTES = int(os.environ.get("MAX_UPLOAD_BYTES", str(5 * 1024 * 1024)))

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("FLASK_SECRET_KEY", "local-dev-only")
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_BYTES


CONVERSION_PRESETS = {
    "markdown": "Convert the document into clean, well-structured Markdown.",
    "html": "Convert the document into semantic HTML. Return only the HTML body content.",
    "plain": "Convert the document into clean plain text with readable spacing.",
    "json": "Convert the document into structured JSON. Return valid JSON only.",
    "summary": "Convert the document into a concise executive summary with headings and bullets.",
    "csv": "Convert tabular content into CSV. Return CSV only.",
}


TEXT_FILE_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".html",
    ".htm",
    ".xml",
    ".log",
}


def fetch_models():
    response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=8)
    response.raise_for_status()
    data = response.json()
    return sorted(model["name"] for model in data.get("models", []) if model.get("name"))


def extract_text(upload):
    if not upload or not upload.filename:
        return ""

    filename = upload.filename
    suffix = Path(filename).suffix.lower()
    raw = upload.read()

    if suffix in TEXT_FILE_SUFFIXES:
        return raw.decode("utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("Install optional dependency pypdf to upload PDF files.") from exc

        reader = PdfReader(BytesIO(raw))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages.append(f"Page {index}\n{page_text}")
        return "\n\n".join(pages).strip()

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Install optional dependency python-docx to upload DOCX files.") from exc

        document = Document(BytesIO(raw))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        tables = []
        for table_index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                tables.append(f"Table {table_index}\n" + "\n".join(rows))
        return "\n\n".join(paragraphs + tables).strip()

    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("Install optional dependency openpyxl to upload XLSX files.") from exc

        workbook = load_workbook(BytesIO(raw), data_only=True, read_only=True)
        sheets = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append("\t".join(values).rstrip())
            if rows:
                sheets.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
        workbook.close()
        return "\n\n".join(sheets).strip()

    return raw.decode("utf-8", errors="replace")


def build_prompt(source_text, conversion_type, custom_instruction):
    preset = CONVERSION_PRESETS.get(conversion_type, CONVERSION_PRESETS["markdown"])
    instruction = custom_instruction.strip() or preset
    return f"""You are a careful document conversion engine.

Task:
{instruction}

Rules:
- Preserve meaning and important details.
- Do not invent facts.
- Return only the converted document, with no chatty introduction.

Source document:
---
{source_text}
---"""


def generate_conversion(model, prompt):
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.1,
        },
    }
    response = requests.post(f"{OLLAMA_BASE_URL}/api/generate", json=payload, timeout=OLLAMA_TIMEOUT)
    response.raise_for_status()
    data = response.json()
    return data.get("response", "").strip()


@app.route("/", methods=["GET", "POST"])
def index():
    result = ""
    source_text = ""
    selected_model = ""
    conversion_type = request.form.get("conversion_type", "markdown")
    custom_instruction = request.form.get("custom_instruction", "")

    try:
        models = fetch_models()
    except requests.RequestException as exc:
        models = []
        flash(f"Could not reach Ollama at {OLLAMA_BASE_URL}: {exc}", "error")

    if request.method == "POST":
        selected_model = request.form.get("model", "")
        source_text = request.form.get("source_text", "").strip()

        try:
            uploaded_text = extract_text(request.files.get("document"))
            if uploaded_text:
                source_text = uploaded_text
        except RuntimeError as exc:
            flash(str(exc), "error")
        except Exception as exc:
            flash(f"Could not read uploaded file: {exc}", "error")

        if not selected_model:
            flash("Choose a model from the Ollama server.", "error")
        elif not source_text:
            flash("Paste text or upload a document to convert.", "error")
        else:
            try:
                prompt = build_prompt(source_text, conversion_type, custom_instruction)
                result = generate_conversion(selected_model, prompt)
            except requests.RequestException as exc:
                flash(f"Ollama generation failed: {exc}", "error")
            except json.JSONDecodeError:
                flash("Ollama returned a response that was not valid JSON.", "error")

    if not selected_model and models:
        selected_model = models[0]

    return render_template(
        "index.html",
        conversion_presets=CONVERSION_PRESETS,
        conversion_type=conversion_type,
        custom_instruction=custom_instruction,
        models=models,
        ollama_base_url=OLLAMA_BASE_URL,
        ollama_timeout=OLLAMA_TIMEOUT,
        result=result,
        selected_model=selected_model,
        source_text=source_text,
    )


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=int(os.environ.get("FLASK_PORT", "5000")), debug=True)
